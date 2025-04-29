import io
import base64
from typing import Dict, List

from docx import Document

from .base import FileProcessor


class DocxProcessor(FileProcessor):
    @staticmethod
    def extract_images(doc: Document) -> List[Dict]:
        images = []
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image_data = rel.target_part.blob
                image_format = rel.target_ref.split('.')[-1]
                image_b64 = base64.b64encode(image_data).decode("utf-8")
                images.append({
                    "format": image_format,
                    "data": image_b64
                })
        return images

    async def process(self, content: bytes, filename: str) -> Dict:
        try:
            doc = Document(io.BytesIO(content))
            pages = []
            current_page = {"text": [], "tables": [], "images": []}

            # Process paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    current_page["text"].append(para.text)
                # If total characters exceed 500, creates a new page
                if sum(len(t) for t in current_page["text"]) > 500:
                    pages.append({
                        "page_number": len(pages) + 1,
                        "text": "\n".join(current_page["text"]),
                        "tables": current_page["tables"],
                        "images": current_page["images"]
                    })
                    current_page = {"text": [], "tables": [], "images": []}

            # Process tables
            for table in doc.tables:
                table_data = [[cell.text for cell in row.cells] for row in table.rows]
                current_page["tables"].append(table_data)

            # Process images
            current_page["images"].extend(self.extract_images(doc))

            # Add remaining content as the last page
            if current_page["text"] or current_page["tables"] or current_page["images"]:
                pages.append({
                    "page_number": len(pages) + 1,
                    "text": "\n".join(current_page["text"]),
                    "tables": current_page["tables"],
                    "images": current_page["images"]
                })

            return {
                "filename": filename,
                "pages": pages,
                "page_count": len(pages)
            }
        except Exception as e:
            return {"filename": filename, "error": f"Failed to process DOCX: {str(e)}"}